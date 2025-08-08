import glob
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from fpdf import FPDF

DOCS_DIR = Path(__file__).parent


def html_to_text(path: Path) -> str:
    with path.open(encoding="utf-8") as f:
        soup = BeautifulSoup(f, "html.parser")
    # Collapsing whitespace and returning plain text
    return soup.get_text(separator="\n")


def build_word(output: Path) -> None:
    doc = Document()
    doc.add_heading("Documentación de Kezan", level=1)
    for html_file in sorted(DOCS_DIR.glob("*.html")):
        doc.add_heading(html_file.stem, level=2)
        doc.add_paragraph(html_to_text(html_file))
    doc.save(output)


def build_pdf(output: Path) -> None:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for html_file in sorted(DOCS_DIR.glob("*.html")):
        pdf.add_page()
        for line in html_to_text(html_file).splitlines():
            pdf.multi_cell(0, 10, line)
    pdf.output(str(output))


if __name__ == "__main__":
    build_word(DOCS_DIR / "documentacion_completa.docx")
    build_pdf(DOCS_DIR / "documentacion_interna.pdf")
